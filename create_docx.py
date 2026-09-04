import docx
from docx.shared import Pt, Inches

doc = docx.Document()

# Title
title = doc.add_heading('Amoeba AI: Enterprise Quotation', 0)

# Paragraph
doc.add_paragraph('This quotation outlines the pricing for the Amoeba AI Enterprise 10-User Package, designed for small-to-medium teams to integrate seamlessly with your existing data infrastructure.')
doc.add_paragraph('Instead of managing unpredictable API costs, this package provides a predictable flat rate while offering a massive capacity for team interactions.')

# Heading 1
doc.add_heading('10-User Enterprise Package', level=2)

# Table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Billing Cycle'
hdr_cells[1].text = 'Total Quote'
hdr_cells[2].text = 'Capacity Included (Per Month)'
hdr_cells[3].text = 'What does this mean?'

row1 = table.add_row().cells
row1[0].text = 'Monthly Billed'
row1[1].text = '?1,200 / month'
row1[2].text = '16.5 Million Tokens'
row1[3].text = '~16,500 Chat Queries'

row2 = table.add_row().cells
row2[0].text = 'Annually Billed (15% Discount)'
row2[1].text = '?12,240 / year'
row2[2].text = '198 Million Tokens / yr'
row2[3].text = '~198,000 Chat Queries'

doc.add_paragraph('\n')

# Heading 2
doc.add_heading('What is included in this capacity?', level=2)

doc.add_paragraph('Your team of 10 users receives a shared pool of 16,500,000 AI Tokens every single month. To put that into perspective:')

doc.add_paragraph('Daily Usage: Your team can ask the Amoeba AI Assistant up to 550 complex questions every single day.', style='List Bullet')
doc.add_paragraph('Per-User Breakdown: This allows every single user on your team to run roughly 55 detailed reports, data summaries, or navigation queries daily.', style='List Bullet')
doc.add_paragraph('Data Processing: 16.5 Million tokens is enough processing power for the AI to read, analyze, and summarize the equivalent of 22,000 pages of text or data every month.', style='List Bullet')

doc.add_paragraph('\n')

# Heading 3
doc.add_heading('Features Included in the License:', level=2)
doc.add_paragraph('Unlimited Access to the Amoeba AI Web Interface & Dashboard.', style='List Bullet')
doc.add_paragraph('Hybrid AI Routing (Utilizing both ultra-fast and ultra-intelligent models seamlessly).', style='List Bullet')
doc.add_paragraph('Live ERP Database Synchronization.', style='List Bullet')
doc.add_paragraph('Unlimited Interactive Visualizations: The AI generates Pie Charts, Bar Graphs, and Data Tables using lightweight text (JSON), meaning complex data visualizations cost you virtually nothing in tokens.', style='List Bullet')
doc.add_paragraph('Dedicated Server Hosting & Maintenance.', style='List Bullet')
doc.add_paragraph('Real-time Token and Usage Analytics Dashboard for your managers.', style='List Bullet')

doc.add_paragraph('\n')
p = doc.add_paragraph()
p.add_run('Overage Policy: ').bold = True
p.add_run('If your team exceeds the 16.5 Million token limit in a given month, you will simply be billed a micro-transaction rate of ?0.07 per additional 1,000 tokens used, ensuring your business is never interrupted.')

doc.save('Amoeba_AI_Enterprise_Quotation.docx')
